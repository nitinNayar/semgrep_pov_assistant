"""
Local File Output Client for semgrep_pov_assistant.

This module provides functionality to create and save analysis results
as local files instead of using Google Drive APIs.
"""

import os
import json
from typing import Dict, List, Optional, Union, Any
from datetime import datetime
from pathlib import Path
import logging

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

from .utils.logger import get_logger, log_error, log_warning, log_info, log_debug

logger = get_logger()

class LocalFileClient:
    """
    Local file output client for creating and saving analysis results.
    
    This class provides methods for creating local files with the analysis
    results from call transcripts, avoiding Google Drive API issues.
    """
    
    def __init__(self, config: Optional[Dict] = None):
        """
        Initialize the local file client.
        
        Args:
            config: Configuration dictionary with output settings
        """
        self.config = config or {}
        
        # Output settings
        self.output_dir = self.config.get('paths', {}).get('output_dir', 'data/output')
        self.create_subdirectories = True  # Always create subdirectories for local file output
        
        # Ensure output directory exists
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        
        log_info("Local file client initialized successfully")
    
    def create_call_summary_file(self, call_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Create a call summary file locally.
        
        Args:
            call_data: Call analysis data
            filename: Optional filename (if None, auto-generates)
            
        Returns:
            Path to the created file
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"call_summary_{timestamp}.txt"
            
            # Create subdirectory for this analysis if enabled
            if self.create_subdirectories:
                analysis_dir = Path(self.output_dir) / f"analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                analysis_dir.mkdir(parents=True, exist_ok=True)
                file_path = analysis_dir / filename
            else:
                file_path = Path(self.output_dir) / filename
            
            # Build content
            content = self._build_call_summary_content(call_data)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            log_info(f"Created call summary file: {file_path}")
            return str(file_path)
            
        except Exception as e:
            log_error(f"Failed to create call summary file: {e}")
            raise
    
    def create_json_analysis_file(self, call_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Create a JSON analysis file locally.
        
        Args:
            call_data: Call analysis data
            filename: Optional filename (if None, auto-generates)
            
        Returns:
            Path to the created file
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"call_analysis_{timestamp}.json"
            
            # Create subdirectory for this analysis if enabled
            if self.create_subdirectories:
                analysis_dir = Path(self.output_dir) / f"analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                analysis_dir.mkdir(parents=True, exist_ok=True)
                file_path = analysis_dir / filename
            else:
                file_path = Path(self.output_dir) / filename
            
            # Write JSON file
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(call_data, f, indent=2, ensure_ascii=False)
            
            log_info(f"Created JSON analysis file: {file_path}")
            return str(file_path)
            
        except Exception as e:
            log_error(f"Failed to create JSON analysis file: {e}")
            raise
    
    def create_action_items_file(self, action_items: List[Dict[str, str]], filename: Optional[str] = None) -> str:
        """
        Create an action items file locally.
        
        Args:
            action_items: List of action item dictionaries
            filename: Optional filename (if None, auto-generates)
            
        Returns:
            Path to the created file
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"action_items_{timestamp}.txt"
            
            # Create subdirectory for this analysis if enabled
            if self.create_subdirectories:
                analysis_dir = Path(self.output_dir) / f"analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                analysis_dir.mkdir(parents=True, exist_ok=True)
                file_path = analysis_dir / filename
            else:
                file_path = Path(self.output_dir) / filename
            
            # Build content
            content = self._build_action_items_content(action_items)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            log_info(f"Created action items file: {file_path}")
            return str(file_path)
            
        except Exception as e:
            log_error(f"Failed to create action items file: {e}")
            raise
    
    def create_sentiment_analysis_file(self, sentiment_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Create a sentiment analysis file locally.
        
        Args:
            sentiment_data: Sentiment analysis data
            filename: Optional filename (if None, auto-generates)
            
        Returns:
            Path to the created file
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"sentiment_analysis_{timestamp}.txt"
            
            # Create subdirectory for this analysis if enabled
            if self.create_subdirectories:
                analysis_dir = Path(self.output_dir) / f"analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                analysis_dir.mkdir(parents=True, exist_ok=True)
                file_path = analysis_dir / filename
            else:
                file_path = Path(self.output_dir) / filename
            
            # Build content
            content = self._build_sentiment_analysis_content(sentiment_data)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            log_info(f"Created sentiment analysis file: {file_path}")
            return str(file_path)
            
        except Exception as e:
            log_error(f"Failed to create sentiment analysis file: {e}")
            raise
    
    def create_technical_deployment_analysis_file(self, deployment_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Create a technical deployment analysis file locally.
        
        Args:
            deployment_data: Technical deployment analysis data
            filename: Optional custom filename
            
        Returns:
            Path to the created file
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"technical_deployment_details_{timestamp}.txt"
        
        # Ensure the output directory exists
        output_dir = Path(self.output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Create timestamped subdirectory
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        subdir = output_dir / f"analysis_{timestamp}"
        subdir.mkdir(parents=True, exist_ok=True)
        
        # Create the file path
        file_path = subdir / filename
        
        # Build the content
        content = self._build_technical_deployment_analysis_content(deployment_data)
        
        # Write the file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return str(file_path)

    def create_customer_overview_analysis_file(self, customer_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Create a customer overview analysis file locally.
        
        Args:
            customer_data: Customer overview analysis data
            filename: Optional custom filename
            
        Returns:
            Path to the created file
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"customer_overview_{timestamp}.txt"
        
        # Ensure the output directory exists
        output_dir = Path(self.output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Create timestamped subdirectory
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        subdir = output_dir / f"analysis_{timestamp}"
        subdir.mkdir(parents=True, exist_ok=True)
        
        # Create the file path
        file_path = subdir / filename
        
        # Build the content
        content = self._build_customer_overview_analysis_content(customer_data)
        
        # Write the file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return str(file_path)

    def create_word_document(self, content_data: Dict[str, Any], doc_type: str, filename: Optional[str] = None) -> str:
        """
        Create a Word document with proper formatting.
        
        Args:
            content_data: Analysis data to include in the document
            doc_type: Type of document ('call_summary', 'pov_analysis', 'technical_deployment', 'customer_overview')
            filename: Optional custom filename
            
        Returns:
            Path to the created Word document
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{doc_type}_{timestamp}.docx"
        
        # Ensure the output directory exists
        output_dir = Path(self.output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Create timestamped subdirectory
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        subdir = output_dir / f"analysis_{timestamp}"
        subdir.mkdir(parents=True, exist_ok=True)
        
        # Create the file path
        file_path = subdir / filename
        
        # Create Word document
        doc = Document()
        
        # Apply document formatting based on type
        if doc_type == 'call_summary':
            self._build_call_summary_word_doc(doc, content_data)
        elif doc_type == 'pov_analysis':
            self._build_pov_analysis_word_doc(doc, content_data)
        elif doc_type == 'technical_deployment':
            self._build_technical_deployment_word_doc(doc, content_data)
        elif doc_type == 'customer_overview':
            self._build_customer_overview_word_doc(doc, content_data)
        else:
            # Generic document
            self._build_generic_word_doc(doc, content_data, doc_type)
        
        # Save the document
        doc.save(str(file_path))
        
        log_info(f"Created Word document: {file_path}")
        return str(file_path)
    
    def create_pov_analysis_file(self, pov_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Create a POV Win/Loss analysis file locally.
        
        Args:
            pov_data: POV analysis data
            filename: Optional filename (if None, auto-generates)
            
        Returns:
            Path to the created file
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"pov_analysis_{timestamp}.txt"
            
            # Create subdirectory for this analysis if enabled
            if self.create_subdirectories:
                analysis_dir = Path(self.output_dir) / f"analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                analysis_dir.mkdir(parents=True, exist_ok=True)
                file_path = analysis_dir / filename
            else:
                file_path = Path(self.output_dir) / filename
            
            # Build content
            content = self._build_pov_analysis_content(pov_data)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            log_info(f"Created POV analysis file: {file_path}")
            return str(file_path)
            
        except Exception as e:
            log_error(f"Failed to create POV analysis file: {e}")
            raise
    
    def _build_call_summary_content(self, call_data: Dict[str, Any]) -> str:
        """
        Build content for a call summary file.
        
        Args:
            call_data: Call analysis data
            
        Returns:
            Formatted content string
        """
        metadata = call_data.get('metadata', {})
        analysis = call_data.get('analysis', {})
        action_items = call_data.get('action_items', [])
        sentiment_analysis = call_data.get('sentiment_analysis', {})
        
        content = f"""# Call Summary - Semgrep POV Assistant

## Call Details
- Date: {metadata.get('extracted_date', 'Unknown')}
- File: {metadata.get('filename', 'Unknown')}
- Duration: {metadata.get('duration', 'Unknown')}
- Participants: {', '.join(analysis.get('call_overview', {}).get('Participants', []))}

## Executive Summary
{analysis.get('executive_summary', 'No summary available')}

## Key Discussion Points
{self._format_list(analysis.get('key_discussion_points', []))}

## Business Context
{self._format_dict(analysis.get('business_context', {}))}

## Action Items
{self._format_action_items(action_items)}

## Sentiment Analysis
{self._format_sentiment_analysis(sentiment_analysis)}

## Next Steps
{self._format_list(analysis.get('next_steps', []))}

Generated by Semgrep POV Assistant
Timestamp: {datetime.now().isoformat()}
"""
        
        return content
    
    def _build_action_items_content(self, action_items: List[Dict[str, str]]) -> str:
        """
        Build content for action items file.
        
        Args:
            action_items: List of action item dictionaries
            
        Returns:
            Formatted content string
        """
        if not action_items:
            return "No action items identified."
        
        content = f"""# Action Items - Semgrep POV Assistant

## Summary
Total Action Items: {len(action_items)}

## Action Items List
{self._format_action_items(action_items)}

Generated by Semgrep POV Assistant
Timestamp: {datetime.now().isoformat()}
"""
        
        return content
    
    def _build_sentiment_analysis_content(self, sentiment_data: Dict[str, Any]) -> str:
        """
        Build content for sentiment analysis file.
        
        Args:
            sentiment_data: Sentiment analysis data
            
        Returns:
            Formatted content string
        """
        content = []
        content.append("# Sentiment Analysis - Semgrep POV Assistant")
        content.append("")
        
        # Overall sentiment
        overall_sentiment = sentiment_data.get('overall_sentiment', 'Unknown')
        confidence = sentiment_data.get('confidence', 0)
        content.append(f"## Overall Sentiment: {overall_sentiment}")
        content.append(f"**Confidence**: {confidence}%")
        content.append("")
        
        # Key indicators
        key_indicators = sentiment_data.get('key_indicators', [])
        if key_indicators:
            content.append("## Key Sentiment Indicators")
            for indicator in key_indicators:
                content.append(f"- {indicator}")
            content.append("")
        
        # Engagement level
        engagement_level = sentiment_data.get('engagement_level', 'Unknown')
        content.append(f"## Engagement Level: {engagement_level}")
        content.append("")
        
        # Sentiment changes
        sentiment_changes = sentiment_data.get('sentiment_changes', '')
        if sentiment_changes:
            content.append("## Sentiment Changes Throughout Call")
            content.append(sentiment_changes)
            content.append("")
        
        # Recommendations
        recommendations = sentiment_data.get('recommendations', [])
        if recommendations:
            content.append("## Recommendations")
            for recommendation in recommendations:
                content.append(f"- {recommendation}")
            content.append("")
        
        return "\n".join(content)
    
    def _build_pov_analysis_content(self, pov_data: Dict[str, Any]) -> str:
        """
        Build content for POV analysis file.
        
        Args:
            pov_data: POV analysis data
            
        Returns:
            Formatted content string
        """
        content = []
        content.append("# POV Win/Loss Analysis - Semgrep POV Assistant")
        content.append("")
        
        # Win probability
        win_probability = pov_data.get('win_probability', 0)
        reasoning = pov_data.get('probability_reasoning', 'No reasoning provided')
        content.append(f"## Win Probability: {win_probability}%")
        content.append(f"**Reasoning**: {reasoning}")
        content.append("")
        
        # Key positive factors
        positive_factors = pov_data.get('key_positive_factors', [])
        if positive_factors:
            content.append("## Key Positive Factors")
            for factor in positive_factors:
                content.append(f"✅ {factor}")
            content.append("")
        
        # Key risks
        risks = pov_data.get('key_risks', [])
        if risks:
            content.append("## Key Risks")
            for risk in risks:
                content.append(f"⚠️  **{risk.get('risk', 'Unknown risk')}**")
                content.append(f"   - **Severity**: {risk.get('severity', 'Unknown')}")
                content.append(f"   - **Time Open**: {risk.get('time_open', 'Unknown')}")
                content.append(f"   - **Mitigation**: {risk.get('mitigation', 'No mitigation provided')}")
                content.append("")
        
        # Technical win strategy
        tech_strategy = pov_data.get('technical_win_strategy', {})
        if tech_strategy:
            content.append("## Technical Win Strategy")
            
            # Unresolved technical questions
            unresolved_questions = tech_strategy.get('unresolved_technical_questions', [])
            if unresolved_questions:
                content.append("### 📝 Unresolved Technical Questions")
                for question in unresolved_questions:
                    content.append(f"- {question}")
                content.append("")
            
            # Recommended demonstrations
            recommended_demos = tech_strategy.get('recommended_demonstrations', [])
            if recommended_demos:
                content.append("### 🎯 Recommended Demonstrations")
                for demo in recommended_demos:
                    content.append(f"- {demo}")
                content.append("")
            
            # Competitive advantages
            competitive_advantages = tech_strategy.get('competitive_advantages', [])
            if competitive_advantages:
                content.append("### 🏆 Competitive Advantages")
                for advantage in competitive_advantages:
                    content.append(f"- {advantage}")
                content.append("")
        
        # Next steps
        next_steps = pov_data.get('next_steps', [])
        if next_steps:
            content.append("## Recommended Next Steps")
            for step in next_steps:
                content.append(f"📋 {step}")
            content.append("")
        
        # Key transcript snippets
        snippets = pov_data.get('key_transcript_snippets', [])
        if snippets:
            content.append("## Key Transcript Snippets")
            for snippet in snippets:
                content.append(f"### 📄 {snippet.get('call', 'Unknown call')}")
                content.append(f"**Context**: {snippet.get('context', 'No context')}")
                content.append(f"**Quote**: \"{snippet.get('quote', 'No quote')}\"")
                content.append(f"**Significance**: {snippet.get('significance', 'No significance provided')}")
                content.append("")
        
        # Analysis metadata
        if pov_data.get('fallback_analysis'):
            content.append("---")
            content.append("*Note: This analysis was generated using fallback methods due to insufficient data or analysis errors.*")
        
        return "\n".join(content)
    
    def _build_technical_deployment_analysis_content(self, deployment_data: Dict[str, Any]) -> str:
        """
        Build content for technical deployment analysis file.
        
        Args:
            deployment_data: Technical deployment analysis data
            
        Returns:
            Formatted content string
        """
        content = []
        content.append("# Technical Deployment Details - Semgrep POV Assistant")
        content.append("")
        
        # SCM Platform
        scm = deployment_data.get('scm_platform', {})
        if scm:
            content.append("## 📁 Source Code Management (SCM)")
            content.append(f"**Platform**: {scm.get('platform', 'Unknown')}")
            content.append(f"**Deployment Type**: {scm.get('deployment_type', 'Unknown')}")
            content.append(f"**Details**: {scm.get('details', 'No details provided')}")
            if scm.get('evidence'):
                content.append(f"**Evidence**: {scm.get('evidence')}")
            content.append("")
        
        # CI Pipelines
        ci = deployment_data.get('ci_pipelines', {})
        if ci:
            content.append("## 🔄 Continuous Integration (CI) Pipelines")
            content.append(f"**Primary CI**: {ci.get('primary_ci', 'Unknown')}")
            additional_ci = ci.get('additional_ci', [])
            if additional_ci:
                content.append(f"**Additional CI**: {', '.join(additional_ci)}")
            content.append(f"**Details**: {ci.get('details', 'No details provided')}")
            if ci.get('evidence'):
                content.append(f"**Evidence**: {ci.get('evidence')}")
            content.append("")
        
        # Programming Languages
        languages = deployment_data.get('programming_languages', {})
        if languages:
            content.append("## 💻 Programming Languages")
            primary_langs = languages.get('primary_languages', [])
            if primary_langs:
                content.append(f"**Primary Languages**: {', '.join(primary_langs)}")
            poc_langs = languages.get('poc_focus_languages', [])
            if poc_langs:
                content.append(f"**POC Focus Languages**: {', '.join(poc_langs)}")
            content.append(f"**Details**: {languages.get('details', 'No details provided')}")
            if languages.get('evidence'):
                content.append(f"**Evidence**: {languages.get('evidence')}")
            content.append("")
        
        # Integrations
        integrations = deployment_data.get('integrations', {})
        if integrations:
            content.append("## 🔗 Integrations")
            interested = integrations.get('interested_integrations', [])
            if interested:
                content.append(f"**Interested In**: {', '.join(interested)}")
            current = integrations.get('current_integrations', [])
            if current:
                content.append(f"**Current Integrations**: {', '.join(current)}")
            content.append(f"**Details**: {integrations.get('details', 'No details provided')}")
            if integrations.get('evidence'):
                content.append(f"**Evidence**: {integrations.get('evidence')}")
            content.append("")
        
        # Supply Chain Security
        sca = deployment_data.get('supply_chain_security', {})
        if sca:
            content.append("## 📦 Supply Chain Security")
            sca_langs = sca.get('languages_tested', [])
            if sca_langs:
                content.append(f"**Languages Tested**: {', '.join(sca_langs)}")
            package_mgrs = sca.get('package_managers', [])
            if package_mgrs:
                content.append(f"**Package Managers**: {', '.join(package_mgrs)}")
            content.append(f"**Details**: {sca.get('details', 'No details provided')}")
            if sca.get('evidence'):
                content.append(f"**Evidence**: {sca.get('evidence')}")
            content.append("")
        
        # Current Security Tools
        security_tools = deployment_data.get('current_security_tools', {})
        if security_tools:
            content.append("## 🛡️ Current Security Tools")
            content.append(f"**SAST**: {security_tools.get('sast', 'Unknown')}")
            content.append(f"**DAST**: {security_tools.get('dast', 'Unknown')}")
            content.append(f"**SCA**: {security_tools.get('sca', 'Unknown')}")
            content.append(f"**Secrets Detection**: {security_tools.get('secrets_detection', 'Unknown')}")
            content.append(f"**ASPM**: {security_tools.get('aspm', 'Unknown')}")
            content.append(f"**Details**: {security_tools.get('details', 'No details provided')}")
            if security_tools.get('evidence'):
                content.append(f"**Evidence**: {security_tools.get('evidence')}")
            content.append("")
        
        # IDE Environment
        ide = deployment_data.get('ide_environment', {})
        if ide:
            content.append("## 💻 IDE Environment")
            content.append(f"**Primary IDE**: {ide.get('primary_ide', 'Unknown')}")
            additional_tools = ide.get('additional_tools', [])
            if additional_tools:
                content.append(f"**Additional Tools**: {', '.join(additional_tools)}")
            content.append(f"**Details**: {ide.get('details', 'No details provided')}")
            if ide.get('evidence'):
                content.append(f"**Evidence**: {ide.get('evidence')}")
            content.append("")
        
        # Additional Technical Details
        additional_details = deployment_data.get('additional_technical_details', [])
        if additional_details:
            content.append("## 📋 Additional Technical Details")
            for detail in additional_details:
                content.append(f"• {detail}")
            content.append("")
        
        # Deployment Complexity
        complexity = deployment_data.get('deployment_complexity', 'Unknown')
        content.append(f"## 📊 Deployment Complexity")
        content.append(f"**Level**: {complexity}")
        content.append("")
        
        # Migration Considerations
        migration = deployment_data.get('migration_considerations', [])
        if migration:
            content.append("## 🔄 Migration Considerations")
            for consideration in migration:
                content.append(f"• {consideration}")
            content.append("")
        
        # Technical Risks
        risks = deployment_data.get('technical_risks', [])
        if risks:
            content.append("## ⚠️ Technical Risks")
            for risk in risks:
                content.append(f"• {risk}")
            content.append("")
        
        # Recommendations
        recommendations = deployment_data.get('recommendations', [])
        if recommendations:
            content.append("## 💡 Technical Recommendations")
            for rec in recommendations:
                content.append(f"• {rec}")
            content.append("")
        
        # Add note if this was a fallback analysis
        if deployment_data.get('fallback_analysis'):
            content.append("---")
            content.append("*Note: This analysis was generated using fallback methods due to insufficient data or analysis errors.*")
        
        return "\n".join(content)
    
    def _build_customer_overview_analysis_content(self, customer_data: Dict[str, Any]) -> str:
        """
        Build content for customer overview analysis file.
        
        Args:
            customer_data: Customer overview analysis data
            
        Returns:
            Formatted content string
        """
        content = []
        content.append("# Customer Overview Analysis - Semgrep POV Assistant")
        content.append("")
        
        # Current State
        current_state = customer_data.get('current_state', {})
        if current_state:
            content.append("## 📊 CURRENT STATE")
            content.append(f"**SAST Tooling**: {current_state.get('sast_tooling', 'Unknown')}")
            content.append(f"**SCA Tooling**: {current_state.get('sca_tooling', 'Unknown')}")
            content.append(f"**Secrets Detection**: {current_state.get('secrets_detection', 'Unknown')}")
            
            challenges = current_state.get('overall_challenges', [])
            if challenges:
                content.append("\n**Key Challenges**:")
                for challenge in challenges:
                    content.append(f"• {challenge}")
            
            evidence = current_state.get('evidence', '')
            if evidence:
                content.append(f"\n**Evidence**: {evidence}")
            content.append("")
        
        # Negative Consequences
        negative_consequences = customer_data.get('negative_consequences', {})
        if negative_consequences:
            content.append("## ⚠️ NEGATIVE CONSEQUENCES")
            
            operational_impact = negative_consequences.get('operational_impact', [])
            if operational_impact:
                content.append("### Operational Impact")
                for impact in operational_impact:
                    content.append(f"• {impact}")
                content.append("")
            
            business_impact = negative_consequences.get('business_impact', [])
            if business_impact:
                content.append("### Business Impact")
                for impact in business_impact:
                    content.append(f"• {impact}")
                content.append("")
            
            security_risk = negative_consequences.get('security_risk', [])
            if security_risk:
                content.append("### Security Risk")
                for risk in security_risk:
                    content.append(f"• {risk}")
                content.append("")
            
            evidence = negative_consequences.get('evidence', '')
            if evidence:
                content.append(f"**Evidence**: {evidence}")
            content.append("")
        
        # Desired Future State
        desired_future_state = customer_data.get('desired_future_state', {})
        if desired_future_state:
            content.append("## 🎯 DESIRED FUTURE STATE")
            
            operational_goals = desired_future_state.get('operational_goals', [])
            if operational_goals:
                content.append("### Operational Goals")
                for goal in operational_goals:
                    content.append(f"• {goal}")
                content.append("")
            
            security_goals = desired_future_state.get('security_goals', [])
            if security_goals:
                content.append("### Security Goals")
                for goal in security_goals:
                    content.append(f"• {goal}")
                content.append("")
            
            business_goals = desired_future_state.get('business_goals', [])
            if business_goals:
                content.append("### Business Goals")
                for goal in business_goals:
                    content.append(f"• {goal}")
                content.append("")
            
            evidence = desired_future_state.get('evidence', '')
            if evidence:
                content.append(f"**Evidence**: {evidence}")
            content.append("")
        
        # Key Semgrep Capabilities
        key_capabilities = customer_data.get('key_semgrep_capabilities', {})
        if key_capabilities:
            content.append("## 🔧 KEY SEMGREP CAPABILITIES")
            
            sast_capabilities = key_capabilities.get('sast_capabilities', [])
            if sast_capabilities:
                content.append("### SAST Capabilities")
                for capability in sast_capabilities:
                    content.append(f"• {capability}")
                content.append("")
            
            sca_capabilities = key_capabilities.get('sca_capabilities', [])
            if sca_capabilities:
                content.append("### SCA Capabilities")
                for capability in sca_capabilities:
                    content.append(f"• {capability}")
                content.append("")
            
            secrets_capabilities = key_capabilities.get('secrets_capabilities', [])
            if secrets_capabilities:
                content.append("### Secrets Detection Capabilities")
                for capability in secrets_capabilities:
                    content.append(f"• {capability}")
                content.append("")
            
            integration_capabilities = key_capabilities.get('integration_capabilities', [])
            if integration_capabilities:
                content.append("### Integration Capabilities")
                for capability in integration_capabilities:
                    content.append(f"• {capability}")
                content.append("")
            
            evidence = key_capabilities.get('evidence', '')
            if evidence:
                content.append(f"**Evidence**: {evidence}")
            content.append("")
        
        # POV Strategy
        pov_strategy = customer_data.get('pov_strategy', {})
        if pov_strategy:
            content.append("## 📋 POV STRATEGY")
            
            primary_focus = pov_strategy.get('primary_focus_areas', [])
            if primary_focus:
                content.append("### Primary Focus Areas")
                for focus in primary_focus:
                    content.append(f"• {focus}")
                content.append("")
            
            demo_priorities = pov_strategy.get('demonstration_priorities', [])
            if demo_priorities:
                content.append("### Demonstration Priorities")
                for priority in demo_priorities:
                    content.append(f"• {priority}")
                content.append("")
            
            success_metrics = pov_strategy.get('success_metrics', [])
            if success_metrics:
                content.append("### Success Metrics")
                for metric in success_metrics:
                    content.append(f"• {metric}")
                content.append("")
            
            evidence = pov_strategy.get('evidence', '')
            if evidence:
                content.append(f"**Evidence**: {evidence}")
            content.append("")
        
        # Key Transcript Snippets
        snippets = customer_data.get('key_transcript_snippets', [])
        if snippets:
            content.append("## 💬 KEY TRANSCRIPT SNIPPETS")
            for snippet in snippets:
                content.append(f"### 📄 {snippet.get('call', 'Unknown call')}")
                content.append(f"**Context**: {snippet.get('context', 'No context')}")
                content.append(f"**Quote**: \"{snippet.get('quote', 'No quote')}\"")
                content.append(f"**Significance**: {snippet.get('significance', 'No significance provided')}")
                content.append("")
        
        # Add note if this was a fallback analysis
        if customer_data.get('fallback_analysis'):
            content.append("---")
            content.append("*Note: This analysis was generated using fallback methods due to insufficient data or analysis errors.*")
        
        return "\n".join(content)
    
    def _format_action_items(self, action_items: List[Dict[str, str]]) -> str:
        """Format action items for display."""
        if not action_items:
            return "No action items identified."
        
        formatted_items = []
        for i, item in enumerate(action_items, 1):
            formatted_item = f"""
### Action Item {i}
- **Action**: {item.get('action', 'No action specified')}
- **Owner**: {item.get('owner', 'Not assigned')}
- **Due Date**: {item.get('due_date', 'Not specified')}
- **Priority**: {item.get('priority', 'Medium')}
- **Context**: {item.get('context', 'No context provided')}
"""
            formatted_items.append(formatted_item)
        
        return '\n'.join(formatted_items)
    
    def _format_sentiment_analysis(self, sentiment_data: Dict[str, Any]) -> str:
        """Format sentiment analysis for display."""
        if not sentiment_data:
            return "No sentiment analysis available."
        
        return f"""
- Overall Sentiment: {sentiment_data.get('overall_sentiment', 'Unknown')}
- Confidence: {sentiment_data.get('confidence', 'Unknown')}%
- Engagement Level: {sentiment_data.get('engagement_level', 'Unknown')}
- Key Indicators: {', '.join(sentiment_data.get('key_indicators', []))}
"""
    
    def _format_list(self, items: List[str]) -> str:
        """Format a list of items."""
        if not items:
            return "None identified."
        
        return '\n'.join([f"- {item}" for item in items])
    
    def _format_dict(self, data: Dict[str, str]) -> str:
        """Format a dictionary of key-value pairs."""
        if not data:
            return "None identified."
        
        return '\n'.join([f"- {key}: {value}" for key, value in data.items()]) 

    def _build_call_summary_word_doc(self, doc: Document, call_data: Dict[str, Any]) -> None:
        """Build a call summary Word document with proper formatting."""
        
        # Title
        title = doc.add_heading('Call Summary - Semgrep POV Assistant', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Call Details section
        doc.add_heading('Call Details', level=1)
        
        metadata = call_data.get('metadata', {})
        analysis = call_data.get('analysis', {})
        
        # Create a table for call details
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Value'
        
        # Add call details
        details = [
            ('Date', metadata.get('extracted_date', 'Unknown')),
            ('File', metadata.get('filename', 'Unknown')),
            ('Duration', metadata.get('duration', 'Unknown')),
            ('Participants', ', '.join(analysis.get('call_overview', {}).get('Participants', [])))
        ]
        
        for field, value in details:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = value
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(analysis.get('executive_summary', 'No summary available'))
        
        # Key Discussion Points
        doc.add_heading('Key Discussion Points', level=1)
        discussion_points = analysis.get('key_discussion_points', [])
        for point in discussion_points:
            doc.add_paragraph(point, style='List Bullet')
        
        # Business Context
        doc.add_heading('Business Context', level=1)
        business_context = analysis.get('business_context', {})
        for key, value in business_context.items():
            p = doc.add_paragraph()
            p.add_run(f'{key}: ').bold = True
            if isinstance(value, list):
                p.add_run(', '.join(value))
            else:
                p.add_run(str(value))
        
        # Action Items
        doc.add_heading('Action Items', level=1)
        action_items = call_data.get('action_items', [])
        for item in action_items:
            p = doc.add_paragraph()
            p.add_run(f"{item.get('action', 'No action specified')}").bold = True
            p.add_run(f" (Owner: {item.get('owner', 'Not assigned')}, Due: {item.get('due_date', 'Not specified')}, Priority: {item.get('priority', 'Medium')})")
        
        # Sentiment Analysis
        doc.add_heading('Sentiment Analysis', level=1)
        sentiment_data = call_data.get('sentiment_analysis', {})
        if sentiment_data:
            p = doc.add_paragraph()
            p.add_run(f"Overall Sentiment: {sentiment_data.get('overall_sentiment', 'Unknown')}").bold = True
            p.add_run(f" (Confidence: {sentiment_data.get('confidence', 'Unknown')}%)")
            
            indicators = sentiment_data.get('key_indicators', [])
            if indicators:
                doc.add_paragraph('Key Indicators:')
                for indicator in indicators:
                    doc.add_paragraph(indicator, style='List Bullet')
        
        # Next Steps
        doc.add_heading('Next Steps', level=1)
        next_steps = analysis.get('next_steps', [])
        for step in next_steps:
            doc.add_paragraph(step, style='List Bullet')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated by Semgrep POV Assistant on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _build_pov_analysis_word_doc(self, doc: Document, pov_data: Dict[str, Any]) -> None:
        """Build a POV analysis Word document with proper formatting."""
        
        # Title
        title = doc.add_heading('POV Win/Loss Analysis - Semgrep POV Assistant', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Win Probability
        doc.add_heading('Win Probability Assessment', level=1)
        win_probability = pov_data.get('win_probability', 0)
        reasoning = pov_data.get('probability_reasoning', 'No reasoning provided')
        
        p = doc.add_paragraph()
        p.add_run(f"Win Probability: {win_probability}%").bold = True
        doc.add_paragraph(f"Reasoning: {reasoning}")
        
        # Key Positive Factors
        doc.add_heading('Key Positive Factors', level=1)
        positive_factors = pov_data.get('key_positive_factors', [])
        for factor in positive_factors:
            doc.add_paragraph(f"✅ {factor}", style='List Bullet')
        
        # Key Risks
        doc.add_heading('Key Risks', level=1)
        risks = pov_data.get('key_risks', [])
        for risk in risks:
            p = doc.add_paragraph()
            p.add_run(f"⚠️ {risk.get('risk', 'Unknown risk')}").bold = True
            p.add_run(f" (Severity: {risk.get('severity', 'Unknown')}, Time Open: {risk.get('time_open', 'Unknown')})")
            doc.add_paragraph(f"Mitigation: {risk.get('mitigation', 'No mitigation provided')}")
        
        # Technical Win Strategy
        doc.add_heading('Technical Win Strategy', level=1)
        tech_strategy = pov_data.get('technical_win_strategy', {})
        
        if tech_strategy:
            # Unresolved Technical Questions
            doc.add_heading('Unresolved Technical Questions', level=2)
            unresolved_questions = tech_strategy.get('unresolved_technical_questions', [])
            for question in unresolved_questions:
                doc.add_paragraph(question, style='List Bullet')
            
            # Recommended Demonstrations
            doc.add_heading('Recommended Demonstrations', level=2)
            recommended_demos = tech_strategy.get('recommended_demonstrations', [])
            for demo in recommended_demos:
                doc.add_paragraph(demo, style='List Bullet')
            
            # Competitive Advantages
            doc.add_heading('Competitive Advantages', level=2)
            competitive_advantages = tech_strategy.get('competitive_advantages', [])
            for advantage in competitive_advantages:
                doc.add_paragraph(advantage, style='List Bullet')
        
        # Next Steps
        doc.add_heading('Recommended Next Steps', level=1)
        next_steps = pov_data.get('next_steps', [])
        for step in next_steps:
            doc.add_paragraph(f"📋 {step}", style='List Bullet')
        
        # Key Transcript Snippets
        doc.add_heading('Key Transcript Snippets', level=1)
        snippets = pov_data.get('key_transcript_snippets', [])
        for snippet in snippets:
            p = doc.add_paragraph()
            p.add_run(f"📄 {snippet.get('call', 'Unknown call')}").bold = True
            doc.add_paragraph(f"Context: {snippet.get('context', 'No context')}")
            doc.add_paragraph(f"Quote: \"{snippet.get('quote', 'No quote')}\"")
            doc.add_paragraph(f"Significance: {snippet.get('significance', 'No significance provided')}")
            doc.add_paragraph()
        
        # Footer
        if pov_data.get('fallback_analysis'):
            doc.add_paragraph()
            note = doc.add_paragraph("Note: This analysis was generated using fallback methods due to insufficient data or analysis errors.")
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated by Semgrep POV Assistant on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _build_technical_deployment_word_doc(self, doc: Document, deployment_data: Dict[str, Any]) -> None:
        """Build a technical deployment analysis Word document with proper formatting."""
        
        # Title
        title = doc.add_heading('Technical Deployment Details - Semgrep POV Assistant', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # SCM Platform
        doc.add_heading('Source Code Management (SCM)', level=1)
        scm = deployment_data.get('scm_platform', {})
        if scm:
            p = doc.add_paragraph()
            p.add_run(f"Platform: {scm.get('platform', 'Unknown')}").bold = True
            p.add_run(f" | Deployment Type: {scm.get('deployment_type', 'Unknown')}")
            doc.add_paragraph(f"Details: {scm.get('details', 'No details provided')}")
            if scm.get('evidence'):
                doc.add_paragraph(f"Evidence: {scm.get('evidence')}")
        
        # CI Pipelines
        doc.add_heading('Continuous Integration (CI) Pipelines', level=1)
        ci = deployment_data.get('ci_pipelines', {})
        if ci:
            p = doc.add_paragraph()
            p.add_run(f"Primary CI: {ci.get('primary_ci', 'Unknown')}").bold = True
            additional_ci = ci.get('additional_ci', [])
            if additional_ci:
                p.add_run(f" | Additional CI: {', '.join(additional_ci)}")
            doc.add_paragraph(f"Details: {ci.get('details', 'No details provided')}")
            if ci.get('evidence'):
                doc.add_paragraph(f"Evidence: {ci.get('evidence')}")
        
        # Programming Languages
        doc.add_heading('Programming Languages', level=1)
        languages = deployment_data.get('programming_languages', {})
        if languages:
            primary_langs = languages.get('primary_languages', [])
            if primary_langs:
                p = doc.add_paragraph()
                p.add_run(f"Primary Languages: {', '.join(primary_langs)}").bold = True
            poc_langs = languages.get('poc_focus_languages', [])
            if poc_langs:
                p = doc.add_paragraph()
                p.add_run(f"POC Focus Languages: {', '.join(poc_langs)}").bold = True
            doc.add_paragraph(f"Details: {languages.get('details', 'No details provided')}")
            if languages.get('evidence'):
                doc.add_paragraph(f"Evidence: {languages.get('evidence')}")
        
        # Integrations
        doc.add_heading('Integrations', level=1)
        integrations = deployment_data.get('integrations', {})
        if integrations:
            interested = integrations.get('interested_integrations', [])
            if interested:
                p = doc.add_paragraph()
                p.add_run(f"Interested In: {', '.join(interested)}").bold = True
            current = integrations.get('current_integrations', [])
            if current:
                p = doc.add_paragraph()
                p.add_run(f"Current Integrations: {', '.join(current)}").bold = True
            doc.add_paragraph(f"Details: {integrations.get('details', 'No details provided')}")
            if integrations.get('evidence'):
                doc.add_paragraph(f"Evidence: {integrations.get('evidence')}")
        
        # Supply Chain Security
        doc.add_heading('Supply Chain Security', level=1)
        sca = deployment_data.get('supply_chain_security', {})
        if sca:
            sca_langs = sca.get('languages_tested', [])
            if sca_langs:
                p = doc.add_paragraph()
                p.add_run(f"Languages Tested: {', '.join(sca_langs)}").bold = True
            package_mgrs = sca.get('package_managers', [])
            if package_mgrs:
                p = doc.add_paragraph()
                p.add_run(f"Package Managers: {', '.join(package_mgrs)}").bold = True
            doc.add_paragraph(f"Details: {sca.get('details', 'No details provided')}")
            if sca.get('evidence'):
                doc.add_paragraph(f"Evidence: {sca.get('evidence')}")
        
        # Current Security Tools
        doc.add_heading('Current Security Tools', level=1)
        security_tools = deployment_data.get('current_security_tools', {})
        if security_tools:
            tools_table = doc.add_table(rows=1, cols=2)
            tools_table.style = 'Table Grid'
            hdr_cells = tools_table.rows[0].cells
            hdr_cells[0].text = 'Tool Type'
            hdr_cells[1].text = 'Current Tool'
            
            tool_types = ['SAST', 'DAST', 'SCA', 'Secrets Detection', 'ASPM']
            for tool_type in tool_types:
                row_cells = tools_table.add_row().cells
                row_cells[0].text = tool_type
                row_cells[1].text = security_tools.get(tool_type.lower().replace(' ', '_'), 'Unknown')
            
            doc.add_paragraph(f"Details: {security_tools.get('details', 'No details provided')}")
            if security_tools.get('evidence'):
                doc.add_paragraph(f"Evidence: {security_tools.get('evidence')}")
        
        # IDE Environment
        doc.add_heading('IDE Environment', level=1)
        ide = deployment_data.get('ide_environment', {})
        if ide:
            p = doc.add_paragraph()
            p.add_run(f"Primary IDE: {ide.get('primary_ide', 'Unknown')}").bold = True
            additional_tools = ide.get('additional_tools', [])
            if additional_tools:
                p.add_run(f" | Additional Tools: {', '.join(additional_tools)}")
            doc.add_paragraph(f"Details: {ide.get('details', 'No details provided')}")
            if ide.get('evidence'):
                doc.add_paragraph(f"Evidence: {ide.get('evidence')}")
        
        # Additional Technical Details
        additional_details = deployment_data.get('additional_technical_details', [])
        if additional_details:
            doc.add_heading('Additional Technical Details', level=1)
            for detail in additional_details:
                doc.add_paragraph(detail, style='List Bullet')
        
        # Deployment Complexity
        doc.add_heading('Deployment Complexity', level=1)
        complexity = deployment_data.get('deployment_complexity', 'Unknown')
        doc.add_paragraph(f"Level: {complexity}")
        
        # Migration Considerations
        migration = deployment_data.get('migration_considerations', [])
        if migration:
            doc.add_heading('Migration Considerations', level=1)
            for consideration in migration:
                doc.add_paragraph(consideration, style='List Bullet')
        
        # Technical Risks
        risks = deployment_data.get('technical_risks', [])
        if risks:
            doc.add_heading('Technical Risks', level=1)
            for risk in risks:
                doc.add_paragraph(risk, style='List Bullet')
        
        # Recommendations
        recommendations = deployment_data.get('recommendations', [])
        if recommendations:
            doc.add_heading('Technical Recommendations', level=1)
            for rec in recommendations:
                doc.add_paragraph(rec, style='List Bullet')
        
        # Footer
        if deployment_data.get('fallback_analysis'):
            doc.add_paragraph()
            note = doc.add_paragraph("Note: This analysis was generated using fallback methods due to insufficient data or analysis errors.")
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated by Semgrep POV Assistant on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _build_customer_overview_word_doc(self, doc: Document, customer_data: Dict[str, Any]) -> None:
        """Build a customer overview analysis Word document with proper formatting."""
        
        # Title
        title = doc.add_heading('Customer Overview Analysis - Semgrep POV Assistant', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Current State
        doc.add_heading('Current State', level=1)
        current_state = customer_data.get('current_state', {})
        if current_state:
            p = doc.add_paragraph()
            p.add_run(f"SAST Tooling: {current_state.get('sast_tooling', 'Unknown')}").bold = True
            p = doc.add_paragraph()
            p.add_run(f"SCA Tooling: {current_state.get('sca_tooling', 'Unknown')}").bold = True
            p = doc.add_paragraph()
            p.add_run(f"Secrets Detection: {current_state.get('secrets_detection', 'Unknown')}").bold = True
            
            challenges = current_state.get('overall_challenges', [])
            if challenges:
                doc.add_heading('Key Challenges', level=2)
                for challenge in challenges:
                    doc.add_paragraph(challenge, style='List Bullet')
            
            evidence = current_state.get('evidence', '')
            if evidence:
                doc.add_paragraph(f"Evidence: {evidence}")
        
        # Negative Consequences
        doc.add_heading('Negative Consequences', level=1)
        negative_consequences = customer_data.get('negative_consequences', {})
        if negative_consequences:
            # Operational Impact
            operational_impact = negative_consequences.get('operational_impact', [])
            if operational_impact:
                doc.add_heading('Operational Impact', level=2)
                for impact in operational_impact:
                    doc.add_paragraph(impact, style='List Bullet')
            
            # Business Impact
            business_impact = negative_consequences.get('business_impact', [])
            if business_impact:
                doc.add_heading('Business Impact', level=2)
                for impact in business_impact:
                    doc.add_paragraph(impact, style='List Bullet')
            
            # Security Risk
            security_risk = negative_consequences.get('security_risk', [])
            if security_risk:
                doc.add_heading('Security Risk', level=2)
                for risk in security_risk:
                    doc.add_paragraph(risk, style='List Bullet')
            
            evidence = negative_consequences.get('evidence', '')
            if evidence:
                doc.add_paragraph(f"Evidence: {evidence}")
        
        # Desired Future State
        doc.add_heading('Desired Future State', level=1)
        desired_future_state = customer_data.get('desired_future_state', {})
        if desired_future_state:
            # Operational Goals
            operational_goals = desired_future_state.get('operational_goals', [])
            if operational_goals:
                doc.add_heading('Operational Goals', level=2)
                for goal in operational_goals:
                    doc.add_paragraph(goal, style='List Bullet')
            
            # Security Goals
            security_goals = desired_future_state.get('security_goals', [])
            if security_goals:
                doc.add_heading('Security Goals', level=2)
                for goal in security_goals:
                    doc.add_paragraph(goal, style='List Bullet')
            
            # Business Goals
            business_goals = desired_future_state.get('business_goals', [])
            if business_goals:
                doc.add_heading('Business Goals', level=2)
                for goal in business_goals:
                    doc.add_paragraph(goal, style='List Bullet')
            
            evidence = desired_future_state.get('evidence', '')
            if evidence:
                doc.add_paragraph(f"Evidence: {evidence}")
        
        # Key Semgrep Capabilities
        doc.add_heading('Key Semgrep Capabilities', level=1)
        key_capabilities = customer_data.get('key_semgrep_capabilities', {})
        if key_capabilities:
            # SAST Capabilities
            sast_capabilities = key_capabilities.get('sast_capabilities', [])
            if sast_capabilities:
                doc.add_heading('SAST Capabilities', level=2)
                for capability in sast_capabilities:
                    doc.add_paragraph(capability, style='List Bullet')
            
            # SCA Capabilities
            sca_capabilities = key_capabilities.get('sca_capabilities', [])
            if sca_capabilities:
                doc.add_heading('SCA Capabilities', level=2)
                for capability in sca_capabilities:
                    doc.add_paragraph(capability, style='List Bullet')
            
            # Secrets Capabilities
            secrets_capabilities = key_capabilities.get('secrets_capabilities', [])
            if secrets_capabilities:
                doc.add_heading('Secrets Detection Capabilities', level=2)
                for capability in secrets_capabilities:
                    doc.add_paragraph(capability, style='List Bullet')
            
            # Integration Capabilities
            integration_capabilities = key_capabilities.get('integration_capabilities', [])
            if integration_capabilities:
                doc.add_heading('Integration Capabilities', level=2)
                for capability in integration_capabilities:
                    doc.add_paragraph(capability, style='List Bullet')
            
            evidence = key_capabilities.get('evidence', '')
            if evidence:
                doc.add_paragraph(f"Evidence: {evidence}")
        
        # POV Strategy
        doc.add_heading('POV Strategy', level=1)
        pov_strategy = customer_data.get('pov_strategy', {})
        if pov_strategy:
            # Primary Focus Areas
            primary_focus = pov_strategy.get('primary_focus_areas', [])
            if primary_focus:
                doc.add_heading('Primary Focus Areas', level=2)
                for focus in primary_focus:
                    doc.add_paragraph(focus, style='List Bullet')
            
            # Demonstration Priorities
            demo_priorities = pov_strategy.get('demonstration_priorities', [])
            if demo_priorities:
                doc.add_heading('Demonstration Priorities', level=2)
                for priority in demo_priorities:
                    doc.add_paragraph(priority, style='List Bullet')
            
            # Success Metrics
            success_metrics = pov_strategy.get('success_metrics', [])
            if success_metrics:
                doc.add_heading('Success Metrics', level=2)
                for metric in success_metrics:
                    doc.add_paragraph(metric, style='List Bullet')
            
            evidence = pov_strategy.get('evidence', '')
            if evidence:
                doc.add_paragraph(f"Evidence: {evidence}")
        
        # Key Transcript Snippets
        doc.add_heading('Key Transcript Snippets', level=1)
        snippets = customer_data.get('key_transcript_snippets', [])
        for snippet in snippets:
            p = doc.add_paragraph()
            p.add_run(f"📄 {snippet.get('call', 'Unknown call')}").bold = True
            doc.add_paragraph(f"Context: {snippet.get('context', 'No context')}")
            doc.add_paragraph(f"Quote: \"{snippet.get('quote', 'No quote')}\"")
            doc.add_paragraph(f"Significance: {snippet.get('significance', 'No significance provided')}")
            doc.add_paragraph()
        
        # Footer
        if customer_data.get('fallback_analysis'):
            doc.add_paragraph()
            note = doc.add_paragraph("Note: This analysis was generated using fallback methods due to insufficient data or analysis errors.")
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated by Semgrep POV Assistant on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _build_generic_word_doc(self, doc: Document, content_data: Dict[str, Any], doc_type: str) -> None:
        """Build a generic Word document with basic formatting."""
        
        # Title
        title = doc.add_heading(f'{doc_type.replace("_", " ").title()} - Semgrep POV Assistant', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content as paragraphs
        for key, value in content_data.items():
            if isinstance(value, (str, int, float)):
                doc.add_heading(key.replace('_', ' ').title(), level=1)
                doc.add_paragraph(str(value))
            elif isinstance(value, list):
                doc.add_heading(key.replace('_', ' ').title(), level=1)
                for item in value:
                    if isinstance(item, dict):
                        for sub_key, sub_value in item.items():
                            p = doc.add_paragraph()
                            p.add_run(f"{sub_key}: ").bold = True
                            p.add_run(str(sub_value))
                    else:
                        doc.add_paragraph(f"• {item}", style='List Bullet')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated by Semgrep POV Assistant on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER 